from pathlib import Path
from docx import Document


def add_markdown_like(doc: Document, title: str, text: str) -> None:
    doc.add_heading(title, level=1)
    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        if not line:
            doc.add_paragraph("")
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("|") and line.endswith("|"):
            # Keep table rows as plain text lines for fidelity.
            doc.add_paragraph(line)
        else:
            doc.add_paragraph(line)


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    docs_dir = root / "docs"
    exchange_doc = root / "exchange-docs" / "home.md"
    target = docs_dir / "MASTER_DOCUMENTATION.docx"

    doc = Document()
    doc.add_heading("American Airlines API - Consolidated Documentation", level=0)
    doc.add_paragraph(
        "This Word document consolidates all key project documentation "
        "for architecture, lifecycle, prompts, metrics, and resource indexing."
    )

    files = [
        ("Master Index", docs_dir / "MASTER_DOCUMENTATION.md"),
        ("Resource Index", docs_dir / "RESOURCE_INDEX.md"),
        ("Exchange Home", exchange_doc),
        ("HLD and LLD", docs_dir / "HLD_LLD.md"),
        ("Project Lifecycle and Workflows", docs_dir / "PROJECT_LIFECYCLE_AND_WORKFLOWS.md"),
        ("Prompt Execution Ledger", docs_dir / "PROMPT_EXECUTION_LEDGER.md"),
        ("Cursor TDLC Metrics", docs_dir / "CURSOR_TDLC_METRICS_FILLED.md"),
        ("Executive Scorecard", docs_dir / "EXECUTIVE_SCORECARD.md"),
    ]

    for section_title, path in files:
        if path.exists():
            text = path.read_text(encoding="utf-8")
            add_markdown_like(doc, section_title, text)
        else:
            doc.add_heading(section_title, level=1)
            doc.add_paragraph(f"Missing source file: {path}")

    doc.save(target)
    print(target)


if __name__ == "__main__":
    main()
